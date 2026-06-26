import io
import os
import yara
import docx
import openpyxl
import zipfile
import xml.etree.ElementTree as ET
from typing import List, Optional
from pydantic import BaseModel, Field
from pdfminer.high_level import extract_text as pdf_extract_text
from backend.app.agents.base import BaseAgent

class FileScanResult(BaseModel):
    is_safe: bool = Field(True, description="True if the file scanned clean of threats and exploits")
    detected_threats: List[str] = Field(default_factory=list, description="List of identified threat names or indicators")
    explanation: str = Field("No file attachment scanned.", description="Details of the file analysis findings")
    extracted_text: Optional[str] = Field(None, description="The sanitized text extracted from the document")

# Compile YARA rules for local file sandboxing
YARA_RULES_SRC = r"""
rule OfficeMacroTrigger {
    strings:
        $s1 = "Document_Open" ascii wide nocase
        $s2 = "AutoOpen" ascii wide nocase
        $s3 = "AutoExec" ascii wide nocase
        $s4 = "Workbook_Open" ascii wide nocase
        $s5 = "DocumentOpen" ascii wide nocase
        $s6 = "Auto_Open" ascii wide nocase
        $s7 = "VBA" ascii wide
    condition:
        any of them
}

rule ExecutableInDocument {
    strings:
        $mz = { 4D 5A }
    condition:
        $mz at 0 and uint32(uint32(0x3c)) == 0x00004550
}

rule PDFMaliciousTriggers {
    strings:
        $js = "/JavaScript" ascii wide nocase
        $js_short = "/JS" ascii wide nocase
        $openaction = "/OpenAction" ascii wide nocase
        $launch = "/Launch" ascii wide nocase
        $aa = "/AA" ascii wide nocase
    condition:
        any of them
}
"""

try:
    compiled_rules = yara.compile(source=YARA_RULES_SRC)
except Exception:
    compiled_rules = None

class FileSecurityAgent(BaseAgent):
    """
    File Security Screening Agent.
    Runs local sandboxed inspection:
    1. Extension validation & size verification.
    2. Magic bytes / PE header detection.
    3. Office document VBA macro detection & PDF dynamic trigger scans.
    4. Clean text extraction from PDF, DOCX, XLSX, PPTX, and TXT files.
    """
    def __init__(self):
        system_instruction = (
            "You are a file safety classifier. Analyze file metadata and structure details for "
            "potential macros, hidden executable code, scripts, or malicious content. "
            "Return a structured JSON output conforming to the response schema."
        )
        super().__init__(system_instruction=system_instruction, model_name="gemini-1.5-flash")

    def scan(self, file_bytes: Optional[bytes] = None, file_name: Optional[str] = None) -> FileScanResult:
        if not file_bytes:
            return FileScanResult(
                is_safe=True,
                detected_threats=[],
                explanation="No file attachment scanned."
            )

        detected_threats = []
        ext = os.path.splitext(file_name.lower())[1] if file_name else ""
        file_size_mb = len(file_bytes) / (1024 * 1024)

        # 1. Blocked Extensions Check
        blocked_extensions = {".exe", ".dll", ".bat", ".ps1", ".sh", ".msi", ".scr", ".cmd", ".vbs", ".com"}
        if ext in blocked_extensions:
            detected_threats.append(f"MALWARE_EXTENSION: Extension '{ext}' is explicitly blocked by safety policy.")

        # 2. File Size Thresholds Check
        if file_size_mb > 5.0:
            detected_threats.append(f"ABNORMAL_FILE_SIZE: File size ({file_size_mb:.2f} MB) exceeds maximum security threshold of 5MB.")

        # 3. Magic Bytes / Embedded Executable PE Check
        if len(file_bytes) >= 2 and file_bytes[:2] == b"MZ":
            detected_threats.append("Executable Payload (MZ Header Detected)")

        # 4. MIME Type Header Integrity
        if ext == ".pdf":
            if len(file_bytes) < 4 or file_bytes[:4] != b"%PDF":
                detected_threats.append("Mismatched MIME Type: Declared PDF, missing standard header magic bytes")
        elif ext in (".docx", ".xlsx", ".pptx"):
            if len(file_bytes) < 4 or file_bytes[:4] != b"PK\x03\x04":
                detected_threats.append("Mismatched MIME Type: Declared Office Document, missing ZIP container magic bytes")

        # 5. YARA Rule Execution
        if compiled_rules:
            try:
                matches = compiled_rules.match(data=file_bytes)
                for match in matches:
                    detected_threats.append(f"YARA Alert: {match.rule} matched.")
            except Exception as e:
                detected_threats.append(f"SANDBOX_SCANNER_ERROR: YARA scan error ({str(e)}).")

        # 6. Deep Container Inspection (Check for hidden vbaProject.bin in openxml containers)
        if ext in (".docx", ".xlsx", ".pptx") and len(file_bytes) >= 4 and file_bytes[:4] == b"PK\x03\x04":
            try:
                with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                    for name in zf.namelist():
                        if "vbaProject" in name or name.endswith(".bin") and "vba" in name.lower():
                            if not any("OfficeMacroTrigger" in t for t in detected_threats):
                                detected_threats.append("YARA Alert: OfficeMacroTrigger matched. (Container scan)")
            except Exception:
                pass

        # 7. Safe Text Extraction & Sanitization
        extracted_text = ""
        parse_status = "Skipped"

        if len(detected_threats) == 0:
            try:
                if ext == ".pdf":
                    parse_status = "Attempted"
                    pdf_file = io.BytesIO(file_bytes)
                    extracted_text = pdf_extract_text(pdf_file)
                    parse_status = "Successful"
                elif ext == ".docx":
                    parse_status = "Attempted"
                    doc = docx.Document(io.BytesIO(file_bytes))
                    full_text = [p.text for p in doc.paragraphs]
                    for table in doc.tables:
                        for row in table.rows:
                            full_text.extend([cell.text for cell in row.cells])
                    extracted_text = "\n".join(full_text)
                    parse_status = "Successful"
                elif ext == ".xlsx":
                    parse_status = "Attempted"
                    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
                    sheet_texts = []
                    for sheet in wb.worksheets:
                        sheet_texts.append(f"Sheet: {sheet.title}")
                        for row in sheet.iter_rows(values_only=True):
                            row_vals = [str(val) for val in row if val is not None]
                            if row_vals:
                                sheet_texts.append(" | ".join(row_vals))
                    extracted_text = "\n".join(sheet_texts)
                    parse_status = "Successful"
                elif ext == ".pptx":
                    parse_status = "Attempted"
                    extracted_text = self._extract_text_from_pptx(file_bytes)
                    parse_status = "Successful"
                elif ext in (".txt", ".json", ".xml", ".csv"):
                    parse_status = "Attempted"
                    extracted_text = file_bytes.decode("utf-8", errors="ignore")
                    parse_status = "Successful"
                else:
                    parse_status = f"Unsupported extension: {ext}"
            except Exception as e:
                parse_status = f"Failed (Error: {str(e)})"
                detected_threats.append(f"DOCUMENT_PARSE_FAILURE: Corrupted document structure or parser error.")

        # Determine final security verdict
        is_safe = len(detected_threats) == 0
        explanation_msg = (
            f"File '{file_name or 'attachment'}' scanned. "
            f"Safety: {'Safe' if is_safe else 'Threats Identified'}. "
            f"Sanitization Parsing: {parse_status}. "
            f"Extracted Content Length: {len(extracted_text)} characters."
        )

        if not is_safe:
            explanation_msg += f" Violations found: {', '.join(detected_threats)}"

        return FileScanResult(
            is_safe=is_safe,
            detected_threats=detected_threats,
            explanation=explanation_msg,
            extracted_text=extracted_text if is_safe else None
        )

    def _extract_text_from_pptx(self, file_bytes: bytes) -> str:
        """Extracts text from a PPTX slide deck without external pptx package dependencies."""
        archive = zipfile.ZipFile(io.BytesIO(file_bytes))
        slide_texts = []
        slide_files = sorted([name for name in archive.namelist() if name.startswith("ppt/slides/slide") and name.endswith(".xml")])
        for slide_file in slide_files:
            xml_content = archive.read(slide_file)
            root = ET.fromstring(xml_content)
            texts = []
            for elem in root.iter():
                if elem.tag.endswith('}t'):
                    if elem.text:
                        texts.append(elem.text)
            if texts:
                slide_texts.append(" ".join(texts))
        return "\n".join(slide_texts)
