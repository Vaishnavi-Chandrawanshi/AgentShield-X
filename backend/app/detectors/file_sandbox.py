import io
import os
import zipfile
import docx
import openpyxl
import xml.etree.ElementTree as ET
from typing import Optional, List
from pdfminer.high_level import extract_text as pdf_extract_text
from backend.app.detectors.base import BaseDetector, DetectorResult

class FileSandboxScanner(BaseDetector):
    """
    File Sandbox Scanner.
    Validates file sizes, checks extension vs header magic bytes,
    inspects openxml container structures, and extracts sanitized text content.
    """
    def __init__(self):
        system_instruction = (
            "You are a file sandbox scanner. Inspect document structural integrity, "
            "file size parameters, container zip archives, and mime-types. "
            "Return a structured JSON response specifying the score (0.00 to 1.00), "
            "confidence (0.00 to 1.00), matched_patterns, explainable reason, and extracted_text."
        )
        super().__init__(system_instruction=system_instruction)

    def detect(
        self,
        prompt: str,
        file_bytes: Optional[bytes] = None,
        file_name: Optional[str] = None
    ) -> DetectorResult:
        if not file_bytes:
            return DetectorResult(score=0.0, confidence=1.0, matched_patterns=[], reason="No file provided to sandbox.")

        return self._local_detect(file_bytes, file_name)

    def _local_detect(self, file_bytes: bytes, file_name: Optional[str]) -> DetectorResult:
        matched_patterns = []

        ext = os.path.splitext(file_name.lower())[1] if file_name else ""
        file_size_mb = len(file_bytes) / (1024 * 1024)

        # 1. File Size Check (> 5MB)
        if file_size_mb > 5.0:
            matched_patterns.append("ABNORMAL_FILE_SIZE")

        # 2. MIME Type Integrity Check
        mismatch = False
        if ext == ".pdf":
            if len(file_bytes) < 4 or file_bytes[:4] != b"%PDF":
                mismatch = True
        elif ext in (".docx", ".xlsx", ".pptx"):
            if len(file_bytes) < 4 or file_bytes[:4] != b"PK\x03\x04":
                mismatch = True

        if mismatch:
            matched_patterns.append("MISMAPPED_MIME_TYPE")

        # 3. Deep container archive scans
        if ext in (".docx", ".xlsx", ".pptx") and len(file_bytes) >= 4 and file_bytes[:4] == b"PK\x03\x04":
            try:
                with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                    for name in zf.namelist():
                        if "vbaProject" in name or (name.endswith(".bin") and "vba" in name.lower()):
                            matched_patterns.append("VBA_MACRO_ARCHIVE")
            except Exception:
                pass

        # 4. Text Extraction and Parser Sanity
        extracted_text = ""
        parse_status = "Successful"
        
        try:
            if ext == ".pdf":
                pdf_file = io.BytesIO(file_bytes)
                extracted_text = pdf_extract_text(pdf_file)
            elif ext == ".docx":
                doc = docx.Document(io.BytesIO(file_bytes))
                full_text = [p.text for p in doc.paragraphs]
                for table in doc.tables:
                    for row in table.rows:
                        full_text.extend([cell.text for cell in row.cells])
                extracted_text = "\n".join(full_text)
            elif ext == ".xlsx":
                wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
                sheet_texts = []
                for sheet in wb.worksheets:
                    sheet_texts.append(f"Sheet: {sheet.title}")
                    for row in sheet.iter_rows(values_only=True):
                        row_vals = [str(val) for val in row if val is not None]
                        if row_vals:
                            sheet_texts.append(" | ".join(row_vals))
                extracted_text = "\n".join(sheet_texts)
            elif ext == ".pptx":
                extracted_text = self._extract_text_from_pptx(file_bytes)
            elif ext in (".txt", ".json", ".xml", ".csv"):
                extracted_text = file_bytes.decode("utf-8", errors="ignore")
            else:
                parse_status = f"Unsupported extension: {ext}"
        except Exception as e:
            parse_status = f"Parser Error ({str(e)})"
            matched_patterns.append("PARSER_STRUCTURE_CORRUPTION")

        from backend.app.detectors.base import compute_deterministic_score
        
        pattern_weights = {
            "ABNORMAL_FILE_SIZE": 0.70,
            "MISMAPPED_MIME_TYPE": 0.80,
            "VBA_MACRO_ARCHIVE": 0.85,
            "PARSER_STRUCTURE_CORRUPTION": 0.60
        }
        
        fake_prompt = f"File: {file_name or 'document'}. Size: {len(file_bytes)} bytes."
        score, confidence, reason = compute_deterministic_score(
            prompt=fake_prompt,
            matched_patterns=matched_patterns,
            detector_severity="MEDIUM",
            default_weight=0.60,
            pattern_weights=pattern_weights,
            detector_name="File Sandbox"
        )
        
        # Format reason to include parse status
        if matched_patterns:
            reason = f"{reason} Extraction: {parse_status}."
        else:
            score = 0.0
            confidence = 1.0
            reason = f"File sandbox check clean. Extraction: {parse_status}. Content Length: {len(extracted_text)} chars."

        return DetectorResult(
            score=score,
            confidence=confidence,
            matched_patterns=matched_patterns,
            reason=reason,
            extracted_text=extracted_text
        )

    def _extract_text_from_pptx(self, file_bytes: bytes) -> str:
        archive = zipfile.ZipFile(io.BytesIO(file_bytes))
        slide_texts = []
        slide_files = sorted([name for name in archive.namelist() if name.startswith("ppt/slides/slide") and name.endswith(".xml")])
        for slide_file in slide_files:
            xml_content = archive.read(slide_file)
            root = ET.fromstring(xml_content)
            texts = []
            for elem in root.iter():
                if elem.tag.endswith('}t'):
                    if elem.text:
                        texts.append(elem.text)
            if texts:
                slide_texts.append(" ".join(texts))
        return "\n".join(slide_texts)
